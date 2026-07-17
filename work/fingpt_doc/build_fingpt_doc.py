from __future__ import annotations

import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Personal Expense Tracker")
WORK = ROOT / "work" / "fingpt_doc"
ASSETS = WORK / "assets"
OUTPUTS = ROOT / "outputs"
REFERENCE = Path(r"E:\NutriGPT AI Augmented Document.docx")
FINAL = OUTPUTS / "FinGPT_AI_Augmented_Project_Documentation.docx"

BLUE = RGBColor(79, 129, 189)
TITLE_BLUE = RGBColor(54, 96, 145)
INK = RGBColor(25, 30, 36)
MONO_BG = "#121820"


def font(name: str, size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    filename = "arialbd.ttf" if bold else ("consola.ttf" if name == "mono" else "arial.ttf")
    return ImageFont.truetype(str(Path(r"C:\Windows\Fonts") / filename), size)


def rounded_box(draw, xy, fill, outline="#95a5b5", radius=15, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered(draw, xy, text, fnt, fill="#18202b"):
    box = draw.textbbox((0, 0), text, font=fnt)
    x = xy[0] + (xy[2] - xy[0] - (box[2] - box[0])) / 2
    y = xy[1] + (xy[3] - xy[1] - (box[3] - box[1])) / 2
    draw.text((x, y), text, font=fnt, fill=fill)


def arrow(draw, start, end, color="#536273", width=4):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    tip = (x2, y2)
    left = (x2 - ux * 16 + px * 8, y2 - uy * 16 + py * 8)
    right = (x2 - ux * 16 - px * 8, y2 - uy * 16 - py * 8)
    draw.polygon([tip, left, right], fill=color)


def make_architecture(path: Path):
    im = Image.new("RGB", (1500, 920), "white")
    d = ImageDraw.Draw(im)
    title = font("sans", 38, True)
    label = font("sans", 25, True)
    small = font("sans", 20)
    d.text((55, 35), "FinGPT Technical Architecture", font=title, fill="#366091")
    boxes = [
        ((80, 150, 420, 355), "React + Vite Client", ["Protected routes", "Charts & forms", "Budgeting assistant"], "#eaf3ff"),
        ((575, 150, 925, 355), "Express REST API", ["Routing/controllers", "JWT authentication", "Validation & errors"], "#eef8f4"),
        ((1080, 120, 1420, 315), "MongoDB", ["Users", "Monthly expenses", "Category budgets"], "#effbef"),
        ((1080, 480, 1420, 675), "Gemini API", ["Auto-categorization", "Monthly insights", "Grounded chat"], "#f5edff"),
        ((575, 515, 925, 720), "Service Layer", ["Expense totals", "Budget usage", "AI context assembly"], "#fff6e8"),
        ((80, 550, 420, 735), "Reports", ["PDFKit monthly report", "Browser download"], "#fff0f2"),
    ]
    for xy, head, lines, fill in boxes:
        rounded_box(d, xy, fill)
        centered(d, (xy[0], xy[1] + 12, xy[2], xy[1] + 65), head, label, "#18202b")
        for i, line in enumerate(lines):
            d.text((xy[0] + 35, xy[1] + 82 + i * 34), f"• {line}", font=small, fill="#344150")
    arrow(d, (420, 250), (575, 250))
    d.text((450, 205), "HTTP/JSON", font=small, fill="#536273")
    arrow(d, (925, 235), (1080, 215))
    arrow(d, (925, 585), (1080, 575))
    arrow(d, (750, 355), (750, 515))
    arrow(d, (575, 620), (420, 640))
    d.text((565, 800), "Deployment: Vercel frontend + Render backend", font=font("sans", 24, True), fill="#366091")
    im.save(path)


def make_er(path: Path):
    im = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(im)
    d.text((55, 35), "FinGPT Data Model", font=font("sans", 38, True), fill="#366091")
    entity_font = font("sans", 25, True)
    field_font = font("sans", 20)
    entities = {
        "User": ((70, 190, 440, 525), ["_id", "firstname", "lastname", "username (unique)", "email (unique)", "password hash"]),
        "Expense Month": ((565, 120, 1000, 600), ["_id", "userId", "monthYear", "monthNumber", "tables[]", "calculations", "  totalIncome", "  totalExpense", "  currentAmount"]),
        "Budget": ((1110, 190, 1450, 525), ["_id", "userId", "category", "amount", "month", "year"]),
        "Transaction Row": ((565, 675, 1000, 850), ["date", "name", "amount", "category"]),
    }
    for name, (xy, fields) in entities.items():
        rounded_box(d, xy, "#f8fbff", "#6b8fb7", 12, 3)
        d.rectangle((xy[0], xy[1], xy[2], xy[1] + 58), fill="#dcecff")
        centered(d, (xy[0], xy[1], xy[2], xy[1] + 58), name, entity_font, "#244f79")
        for i, field in enumerate(fields):
            d.text((xy[0] + 24, xy[1] + 78 + i * 36), field, font=field_font, fill="#273341")
    arrow(d, (440, 350), (565, 350))
    arrow(d, (1000, 350), (1110, 350))
    arrow(d, (780, 600), (780, 675))
    d.text((455, 305), "1 : many", font=field_font, fill="#536273")
    d.text((1015, 305), "1 : many", font=field_font, fill="#536273")
    d.text((795, 620), "embedded rows", font=field_font, fill="#536273")
    im.save(path)


def make_user_flow(path: Path):
    im = Image.new("RGB", (1500, 660), "white")
    d = ImageDraw.Draw(im)
    d.text((55, 30), "Registered User Journey", font=font("sans", 38, True), fill="#366091")
    steps = [
        "Register / Login", "Open Dashboard", "Add Income or Expense", "Review Monthly Ledger",
        "Set Category Budgets", "Generate AI Insights", "Download PDF / Ask Assistant",
    ]
    xs = [45, 250, 455, 660, 865, 1070, 1275]
    for i, (x, step) in enumerate(zip(xs, steps), 1):
        xy = (x, 210, x + 175, 430)
        rounded_box(d, xy, "#f2f7fc", "#75a0c8", 14, 3)
        d.ellipse((x + 57, 235, x + 117, 295), fill="#4f81bd")
        centered(d, (x + 57, 235, x + 117, 295), str(i), font("sans", 26, True), "white")
        words = step.split()
        lines, cur = [], ""
        for word in words:
            if len(cur + " " + word) > 19:
                lines.append(cur)
                cur = word
            else:
                cur = (cur + " " + word).strip()
        if cur:
            lines.append(cur)
        for j, line in enumerate(lines):
            centered(d, (x + 10, 320 + j * 30, x + 165, 350 + j * 30), line, font("sans", 18, j == 0), "#273341")
        if i < len(steps):
            arrow(d, (x + 175, 320), (xs[i], 320), width=3)
    d.text((355, 525), "FinGPT keeps each response grounded in the authenticated user's stored finance data.", font=font("sans", 24, True), fill="#366091")
    im.save(path)


def make_mvc(path: Path):
    im = Image.new("RGB", (1200, 620), "white")
    d = ImageDraw.Draw(im)
    d.text((45, 25), "MVC-Oriented Backend Structure", font=font("sans", 36, True), fill="#366091")
    boxes = [
        ((430, 95, 770, 210), "ROUTES", "HTTP endpoints"),
        ((65, 330, 365, 480), "MODELS", "Mongoose schemas"),
        ((450, 330, 750, 480), "CONTROLLERS", "Request / response"),
        ((835, 330, 1135, 480), "SERVICES", "Business & AI logic"),
    ]
    for xy, head, sub in boxes:
        rounded_box(d, xy, "#eef5fc", "#6f98bf", 15, 3)
        centered(d, (xy[0], xy[1] + 15, xy[2], xy[1] + 70), head, font("sans", 24, True), "#244f79")
        centered(d, (xy[0], xy[1] + 70, xy[2], xy[3] - 10), sub, font("sans", 19), "#344150")
    arrow(d, (600, 210), (600, 330))
    arrow(d, (450, 405), (365, 405))
    arrow(d, (750, 405), (835, 405))
    d.text((395, 530), "MongoDB persistence  |  reusable services  |  thin controllers", font=font("sans", 22, True), fill="#536273")
    im.save(path)


def make_code_panel(path: Path, title: str, lines: list[str], width=1400):
    code_font = font("mono", 20)
    title_font = font("sans", 25, True)
    line_h = 31
    height = 100 + line_h * len(lines) + 35
    im = Image.new("RGB", (width, height), MONO_BG)
    d = ImageDraw.Draw(im)
    d.rectangle((0, 0, width, 68), fill="#1e2834")
    d.ellipse((25, 22, 43, 40), fill="#f25f5c")
    d.ellipse((53, 22, 71, 40), fill="#ffbd2e")
    d.ellipse((81, 22, 99, 40), fill="#28c840")
    d.text((125, 16), title, font=title_font, fill="#d9e4ef")
    for i, line in enumerate(lines, 1):
        y = 88 + (i - 1) * line_h
        d.text((28, y), f"{i:>2}", font=code_font, fill="#607080")
        color = "#8dd7f7" if line.strip().startswith(("const ", "module.exports", "app.use", "router.")) else "#d6deeb"
        d.text((78, y), line, font=code_font, fill=color)
    im.save(path)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.extend([rpr, text_el])
    link.append(run)
    paragraph._p.append(link)


def set_run_font(run, name="Calibri", size=11, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def add_para(doc, text="", style=None, before=0, after=6, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = keep
    if text:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_label_para(doc, label, text, style=None):
    p = add_para(doc, style=style)
    r = p.add_run(label)
    set_run_font(r, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def add_heading(doc, text, level=2):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 2 else 14)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        set_run_font(r, size=14 if level == 2 else 16, bold=True, color=BLUE if level == 2 else TITLE_BLUE)
    return p


def apply_numbering(paragraph, num_id: int, level: int = 0):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.get_or_add_numPr()
    ilvl = numpr.get_or_add_ilvl()
    ilvl.val = level
    numid = numpr.get_or_add_numId()
    numid.val = num_id


def add_bullet(doc, label, text=""):
    p = add_para(doc, after=3)
    apply_numbering(p, 2)
    r = p.add_run(label)
    set_run_font(r, bold=bool(text))
    if text:
        r2 = p.add_run(text)
        set_run_font(r2)
    return p


def add_number(doc, label, text):
    p = add_para(doc, after=4)
    apply_numbering(p, 31)
    r = p.add_run(label)
    set_run_font(r, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def add_figure(doc, path: Path, width=6.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = bool(caption)
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run(caption)
        set_run_font(r, size=9, bold=True, color=RGBColor(88, 96, 105))


def add_page_break(doc):
    doc.add_page_break()


def clear_body(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def build():
    ASSETS.mkdir(parents=True, exist_ok=True)
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    make_architecture(ASSETS / "architecture.png")
    make_er(ASSETS / "er.png")
    make_user_flow(ASSETS / "user-flow.png")
    make_mvc(ASSETS / "mvc.png")
    make_code_panel(ASSETS / "server-tree.png", "server/server", [
        "server.js", "src/app.js", "src/db/connect.js", "src/models/user.model.js",
        "src/models/expense.model.js", "src/models/budget.model.js", "src/controllers/",
        "src/services/expense.service.js", "src/services/budget.service.js", "src/services/ai.service.js",
        "src/routes/", "src/utils/pdfReport.js", ".env.example", "package.json",
    ])
    make_code_panel(ASSETS / "client-tree.png", "client/client/src", [
        "App.js", "theme.css", "utils/api.js", "components/Home/", "components/Track/",
        "components/AddIncomeAndExpense/", "components/BudgetManager/", "components/BudgetAssistant/",
        "components/AiInsights/", "components/Charts/", "components/Login/", "components/Registration/",
        "components/ProtectedRoute/", "components/Header/",
    ])
    make_code_panel(ASSETS / "db-connect.png", "src/db/connect.js", [
        "const mongoose = require('mongoose');",
        "require('dotenv').config();",
        "const db = process.env.MONGO_URI || 'mongodb://127.0.0.1:27017/exp';",
        "mongoose.connect(db, {",
        "  useNewUrlParser: true,",
        "  useUnifiedTopology: true,",
        "}).then(() => console.log('Connection successful'))",
        "  .catch((error) => console.log(`No connection: ${error}`));",
    ])
    make_code_panel(ASSETS / "api-routes.png", "src/app.js - mounted API groups", [
        "app.use('/', authRoutes);",
        "app.use('/', userRoutes);",
        "app.use('/', expenseRoutes);",
        "app.use('/', reportRoutes);",
        "app.use('/ai', aiRoutes);",
        "app.use('/budgets', budgetRoutes);",
    ])

    working = WORK / "FinGPT-working.docx"
    shutil.copy2(REFERENCE, working)
    doc = Document(working)
    clear_body(doc)
    configure_styles(doc)

    title = doc.add_paragraph(style="Heading 1")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("FinGPT")
    set_run_font(r, size=22, bold=True, color=TITLE_BLUE)

    add_heading(doc, "Project Description:")
    add_para(doc, "Welcome to FinGPT - an AI-augmented personal finance tracker that helps users record income and expenses, monitor savings, set category budgets, and understand spending patterns from one structured workspace.")
    add_para(doc, "Built with the MERN stack (MongoDB, Express.js, React, and Node.js), FinGPT combines monthly transaction ledgers, charts, PDF reporting, and a Gemini-powered assistant. The AI layer categorizes transactions, produces grounded monthly insights, and answers budgeting questions using the authenticated user's stored financial data. Amounts are presented in Indian Rupees.")
    add_para(doc, "FinGPT turns day-to-day financial records into clear, practical decisions without spreadsheet clutter.")

    add_heading(doc, "Scenario-Based Case Study:")
    add_para(doc, "A salaried professional wants to control household spending but currently records transactions across notes and spreadsheets. At the end of each month, it is difficult to see how much was earned, where money was spent, whether category limits were exceeded, and what practical action should be taken next.")
    add_heading(doc, "The Solution:")
    add_para(doc, "FinGPT provides a secure monthly finance workspace where the user records income and expenses, reviews calculated totals and charts, sets category budgets, exports a report, and receives AI-generated guidance grounded in the same stored records.")

    add_heading(doc, "How FinGPT Helps Users")
    for label, text in [
        ("Secure Registration and Login - ", "Create an account, sign in with hashed credentials, and use JWT-backed application sessions."),
        ("Monthly Transaction Tracking - ", "Record income and expense rows by month and maintain them with add, edit, and delete actions."),
        ("Automatic Financial Totals - ", "Recalculate total income, total expense, and current savings after each change."),
        ("AI Categorization and Insights - ", "Use Gemini to classify transactions and generate short, data-backed observations."),
        ("Budget Planning - ", "Set monthly limits by category and compare budgeted, spent, remaining, and over-budget amounts."),
        ("Charts and Reports - ", "Review monthly trends and category distribution, then download a PDF summary."),
        ("Grounded Budgeting Assistant - ", "Ask finance questions and receive concise answers based on real stored monthly data."),
    ]:
        add_number(doc, label, text)

    add_heading(doc, "System Requirements:")
    add_para(doc, "The following software and hardware requirements support local development, testing, and deployment of the FinGPT web application.")
    add_heading(doc, "1. Software Requirements")
    for label, text in [
        ("Operating System: ", "Windows 10/11, macOS, or a modern Linux distribution."),
        ("Node.js: ", "A current LTS release compatible with Vite 7 and the server dependencies."),
        ("npm: ", "Used to install and run frontend and backend packages."),
        ("React 18 + Vite: ", "Frontend application and development build tool."),
        ("Express.js 4: ", "REST API and middleware framework."),
        ("MongoDB: ", "Local MongoDB or MongoDB Atlas for users, monthly expenses, and budgets."),
        ("Gemini API key: ", "Required for transaction categorization, insights, and assistant chat."),
        ("Browser: ", "Latest Google Chrome, Microsoft Edge, or Firefox."),
        ("Postman or Insomnia: ", "Recommended for API testing."),
        ("Visual Studio Code and Git: ", "Recommended editor and version-control tools."),
    ]:
        add_bullet(doc, label, text)
    add_heading(doc, "2. Hardware Requirements")
    for label, text in [
        ("Processor: ", "Intel Core i5 / AMD Ryzen 5 or better recommended for smooth local development."),
        ("Memory: ", "8 GB minimum; 16 GB recommended when running the client, server, database, and browser together."),
        ("Storage: ", "At least 1 GB free for source code, dependencies, builds, and local database data."),
        ("Display: ", "1920 x 1080 recommended for dashboard and responsive-layout validation."),
    ]:
        add_bullet(doc, label, text)

    add_heading(doc, "Project Architecture:")
    add_para(doc, "FinGPT uses a modular client-server architecture:")
    add_label_para(doc, "React + Vite client  <==>  Express + Node.js REST API  <==>  MongoDB, with Gemini AI and PDFKit as integrated services.", "")
    add_heading(doc, "Technical Architecture:")
    add_figure(doc, ASSETS / "architecture.png", 6.0, "Figure 1. FinGPT technical architecture")
    add_para(doc, "The frontend communicates with route groups mounted by the Express application. Controllers validate request data and delegate reusable work to services. Mongoose models persist users, monthly expense documents, and category budgets. The AI service uses Gemini for structured categorization and insights plus a conversational budgeting assistant.")

    for title_text, body in [
        ("1. User Interface (UI)", "The React 18 frontend uses Vite, React Router, Bootstrap, styled-components, Recharts, and a Three.js finance scene. Protected routes expose Home, Track, Add Transaction, Budgets, and Assistant pages after authentication."),
        ("2. Web Server", "The Node.js and Express backend parses JSON and URL-encoded requests, enables CORS, connects to MongoDB, and mounts authentication, user, expense, report, AI, and budget routes."),
        ("3. API Gateway / Router", "Route modules map HTTP endpoints to controllers for login and registration, monthly transaction CRUD, PDF generation, budget operations, and AI requests."),
        ("4. Authentication Service", "Registration hashes passwords with bcrypt. Login validates the password and returns a signed JWT plus the user record. The frontend stores the session through cookies for protected navigation."),
        ("5. Database (MongoDB)", "Mongoose schemas store users, monthly expense ledgers with embedded transaction rows, and monthly category budgets. Expense calculations are recomputed after writes."),
        ("6. AI and Reporting Services", "Gemini categorizes transaction rows, creates structured insight cards, and answers chat questions from a compact financial snapshot. PDFKit streams a monthly report to the browser."),
    ]:
        add_heading(doc, title_text)
        add_para(doc, body)

    add_heading(doc, "ER Diagram:")
    add_figure(doc, ASSETS / "er.png", 6.0, "Figure 2. FinGPT document relationships")
    for title_text, text in [
        ("1. User and Expense Relationship", "Type: one-to-many. Each monthly Expense document stores the user's identifier, month/year, transaction tables, and calculated totals."),
        ("2. User and Budget Relationship", "Type: one-to-many. Each Budget document belongs to a user and is uniquely scoped by category, month, and year."),
        ("3. Expense and Transaction Row Relationship", "Type: embedded one-to-many. Each monthly document contains INCOME and EXPENSE tables whose rows store date, name, amount, and optional AI category."),
    ]:
        add_heading(doc, title_text)
        add_para(doc, text)

    add_heading(doc, "Features")
    for label, text in [
        ("Role-aware session flow: ", "Visitors can register or log in; authenticated users access finance routes and can log out."),
        ("Monthly ledger management: ", "Add new month data, filter by month/year, edit or delete individual rows, and delete an entire month."),
        ("Calculated finance summary: ", "Total income, total expense, and current savings remain synchronized with transaction rows."),
        ("Interactive analytics: ", "Recharts renders a monthly trend bar chart and category spending pie chart."),
        ("Spending guardrail: ", "The tracker warns when total expenses exceed half of monthly income."),
        ("AI transaction categorization: ", "Gemini assigns controlled income or expense categories and writes them back to the month."),
        ("AI insight cards: ", "The application returns a short summary and 3-5 severity-tagged insights grounded in real numbers."),
        ("Budget monitoring: ", "Monthly category budgets report spent, remaining, percentage used, and over-budget state."),
        ("PDF report export: ", "A downloadable report lists monthly transactions and totals."),
        ("Deployment readiness: ", "Render and Vercel configuration files support split backend/frontend deployment."),
    ]:
        add_bullet(doc, label, text)

    add_heading(doc, "Roles and Responsibilities")
    add_heading(doc, "1. Visitor")
    add_bullet(doc, "Account access: ", "Register a new account or log in with an existing email and password.")
    add_bullet(doc, "Public navigation: ", "Use authentication pages; protected pages redirect when no JWT cookie is present.")
    add_heading(doc, "2. Registered User")
    for label, text in [
        ("Transaction management: ", "Create, review, edit, and delete monthly income and expense entries."),
        ("Financial review: ", "Inspect totals, trends, category distribution, and spending alerts."),
        ("Budget planning: ", "Set and remove monthly category limits and monitor use."),
        ("AI assistance: ", "Categorize rows, generate insight cards, and ask grounded budgeting questions."),
        ("Reporting: ", "Download a PDF copy of the selected monthly ledger."),
    ]:
        add_bullet(doc, label, text)

    add_heading(doc, "User Flow")
    add_figure(doc, ASSETS / "user-flow.png", 6.0, "Figure 3. Registered user journey")
    add_heading(doc, "MVC Pattern")
    add_figure(doc, ASSETS / "mvc.png", 5.8, "Figure 4. Backend route-controller-service-model flow")
    add_para(doc, "FinGPT follows an MVC-oriented backend with an additional service layer. Models define data, controllers handle request/response responsibilities, routes expose endpoints, and services contain financial calculations, budget aggregation, and Gemini integration.")

    add_heading(doc, "Project Setup and Configuration")
    add_heading(doc, "Creating the Project Folder")
    for label, text in [
        ("1. Clone the repository. ", "Use Git to obtain the project locally."),
        ("2. Open the root folder. ", "The source contains separate client/client and server/server applications."),
        ("3. Install dependencies. ", "Run npm install inside both nested application folders."),
        ("4. Create environment files. ", "Copy each .env.example to .env and fill only local/deployment-specific values."),
    ]:
        add_number(doc, label, text)

    add_heading(doc, "Backend Development:")
    add_heading(doc, "Backend Structure:")
    add_figure(doc, ASSETS / "server-tree.png", 5.5, "Figure 5. Backend project structure")
    for label, text in [
        ("server.js: ", "Loads environment variables, imports the Express app, and listens on PORT (default 5100)."),
        ("src/app.js: ", "Configures middleware, database startup, and route mounting."),
        ("controllers/: ", "Coordinates authentication, expense, budget, AI, user, and report responses."),
        ("services/: ", "Contains reusable calculations, budget aggregation, and Gemini workflows."),
        ("models/: ", "Defines User, Expense, and Budget Mongoose schemas."),
        ("routes/: ", "Declares REST endpoints and maps them to controllers."),
        ("utils/pdfReport.js: ", "Streams the selected month as a PDFKit report."),
    ]:
        add_bullet(doc, label, text)
    add_heading(doc, "API Route Groups")
    add_figure(doc, ASSETS / "api-routes.png", 5.8, "Figure 6. Route groups mounted by the Express application")

    add_heading(doc, "Database Development:")
    add_heading(doc, "1. Configure MongoDB")
    add_para(doc, "Set MONGO_URI in server/server/.env. The supplied example defaults to mongodb://127.0.0.1:27017/exp for local development, while deployment can use MongoDB Atlas.")
    add_figure(doc, ASSETS / "db-connect.png", 5.8, "Figure 7. MongoDB connection module")
    add_heading(doc, "2. Configure Environment Variables")
    for label, text in [
        ("MONGO_URI: ", "MongoDB connection string."),
        ("PORT: ", "Backend port; defaults to 5100."),
        ("GEMINI_API_KEY: ", "Google Gemini credential used only by the server."),
        ("GEMINI_MODEL: ", "Model identifier; the project example uses gemini-2.5-flash."),
        ("JWT_SECRET: ", "Long random key used to sign login tokens."),
        ("VITE_API_URL: ", "Frontend base URL for the deployed or local backend."),
    ]:
        add_bullet(doc, label, text)
    add_heading(doc, "3. Create Schemas")
    add_label_para(doc, "User schema: ", "first name, last name, username, email, and bcrypt password hash.")
    add_label_para(doc, "Expense schema: ", "user/month identity, transaction tables, embedded rows, and derived totals.")
    add_label_para(doc, "Budget schema: ", "user, category, amount, month, and year with a unique compound index.")

    add_heading(doc, "Front-End Development:")
    add_heading(doc, "Structure:")
    add_figure(doc, ASSETS / "client-tree.png", 5.6, "Figure 8. Frontend component structure")
    for label, text in [
        ("App.js: ", "Defines BrowserRouter routes and lazy-loads the global Three.js finance backdrop."),
        ("Home/: ", "Landing dashboard with finance overview, CTA links, and a 3D finance visual."),
        ("Track/: ", "Monthly ledger, charts, row CRUD, AI tools, PDF export, and calculated summaries."),
        ("AddIncomeAndExpense/: ", "Form for monthly income or expense entry with category selection."),
        ("BudgetManager/: ", "Monthly category budget form, progress calculations, and deletion."),
        ("BudgetAssistant/: ", "Gemini chat interface grounded in the signed-in user's data."),
        ("AiInsights/ and Charts/: ", "Reusable AI insight cards and Recharts visualizations."),
        ("utils/api.js: ", "Normalizes VITE_API_URL and produces consistent endpoint URLs."),
    ]:
        add_bullet(doc, label, text)

    add_heading(doc, "Project Execution:")
    add_heading(doc, "Step 1: Start the Backend Server")
    for line in ["cd server/server", "npm install", "copy .env.example .env", "npm start"]:
        p = add_para(doc, line, after=2)
        for run in p.runs:
            set_run_font(run, name="Consolas", size=9, color=RGBColor(40, 55, 70))
    add_para(doc, "The backend listens at http://localhost:5100 unless PORT is changed.")
    add_heading(doc, "Step 2: Start the Frontend")
    for line in ["cd client/client", "npm install", "copy .env.example .env", "npm start"]:
        p = add_para(doc, line, after=2)
        for run in p.runs:
            set_run_font(run, name="Consolas", size=9, color=RGBColor(40, 55, 70))
    add_para(doc, "Vite serves the frontend at http://localhost:3000 by default in the project instructions. Ensure VITE_API_URL points to the backend.")
    add_heading(doc, "Step 3: Deployment")
    add_bullet(doc, "Backend on Render: ", "Use server/server as the root, npm install as the build command, npm start as the start command, and configure the four server environment variables.")
    add_bullet(doc, "Frontend on Vercel: ", "Use client/client as the root, npm run build, dist as the output, and set VITE_API_URL to the Render service URL.")

    add_heading(doc, "Output Screenshots:")
    screenshots = [
        ("Home / Dashboard Page", WORK / "home.png"),
        ("Monthly Tracker Page", WORK / "track.png"),
        ("Add Transaction Page", WORK / "add-transaction.png"),
        ("Budget Manager Page", WORK / "budgets.png"),
        ("AI Budgeting Assistant Page", WORK / "assistant.png"),
    ]
    for i, (caption, path) in enumerate(screenshots):
        add_heading(doc, caption + ":")
        add_figure(doc, path, 5.5, f"Figure {9 + i}. {caption}")

    add_heading(doc, "Code Repository")
    p = add_para(doc, "GitHub: ")
    add_hyperlink(p, "https://github.com/sarit-smartbridge/FinGpt-Finance-Tracker", "https://github.com/sarit-smartbridge/FinGpt-Finance-Tracker")

    doc.core_properties.title = "FinGPT AI Augmented Project Documentation"
    doc.core_properties.subject = "MERN personal finance tracker with Gemini AI"
    doc.core_properties.author = "SmartBridge / SkillWallet"
    doc.save(FINAL)
    print(FINAL)


if __name__ == "__main__":
    build()
